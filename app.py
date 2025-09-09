# app.py
import os
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
import fitz # PyMuPDF
import tempfile
import logging
from datetime import datetime
import uuid
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

app = Flask(__name__, static_folder='static')

# --- Configure Logging ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- CORS Configuration for Production Readiness ---
# This line dynamically sets the allowed origin based on the APP_DOMAIN environment variable,
# which DigitalOcean automatically provides. This is a secure and flexible approach.
app_domain = os.environ.get('APP_DOMAIN', 'localhost')
allowed_origins = [f"https://{app_domain}", "http://localhost"]
CORS(app, resources={
    r"/upload-pdf": {"origins": allowed_origins},
    r"/download-pdf/*": {"origins": allowed_origins},
    r"/download-docx/*": {"origins": allowed_origins}
})

# --- File Upload Configuration ---
app.config['UPLOAD_FOLDER'] = tempfile.gettempdir()
app.config['MAX_CONTENT_LENGTH'] = 64 * 1024 * 1024 # 64 MB file size limit

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
logging.info(f"Upload folder set to: {app.config['UPLOAD_FOLDER']}")

# --- PDF Highlight Extraction Function ---
def extract_highlights_from_pdf(pdf_path):
    """
    Extracts highlighted text along with their page numbers and bounding boxes
    from a PDF document.
    """
    highlights = []
    try:
        doc = fitz.open(pdf_path)
        for page_num, page in enumerate(doc):
            annotations = page.annots()
            for annot in annotations:
                if annot.type[0] == fitz.PDF_ANNOT_HIGHLIGHT:
                    rect = annot.rect
                    if rect and not rect.is_empty:
                        text_in_rect = page.get_text(clip=rect)
                        if text_in_rect:
                            full_highlighted_text = text_in_rect.strip()
                            highlights.append({
                                "text": full_highlighted_text,
                                "page": page_num + 1,
                                "rect": list(rect)
                            })
        doc.close()
        logging.info(f"Successfully extracted {len(highlights)} highlights from {pdf_path}")
    except fitz.FileNotFoundError:
        logging.error(f"PDF file not found at '{pdf_path}'")
        return []
    except Exception as e:
        logging.error(f"Error during PDF processing for extraction: {e}")
        return []
    return highlights

# --- Function to Generate PDF from Highlight Screenshots ---
def generate_pdf_from_highlight_screenshots(original_pdf_path, highlights, output_pdf_path, mode):
    """
    Generates a new PDF document based on the specified mode.
    Mode 'full_page': Takes a screenshot of the entire page with the highlight.
    Mode 'cropped_highlight': Takes a screenshot of just the highlighted area.
    """
    output_doc = fitz.open()
    
    if not highlights:
        page = output_doc.new_page()
        page.insert_text((50, 50), "No highlights found in the uploaded document.", fontsize=12)
        logging.info("Generated PDF with 'No highlights found' message.")
    else:
        try:
            original_doc = fitz.open(original_pdf_path)

            if mode == 'full_page':
                unique_original_page_nums = sorted(list(set(item['page'] for item in highlights)))
                for original_page_num in unique_original_page_nums:
                    original_page = original_doc.load_page(original_page_num - 1)
                    matrix = fitz.Matrix(2, 2) 
                    pixmap = original_page.get_pixmap(matrix=matrix)
                    
                    new_page = output_doc.new_page(width=595, height=842)
                    y_offset = 50
                    margin = 50
                    
                    title_text = f"Original Page {original_page_num} (with highlights):"
                    new_page.insert_text((margin, y_offset), title_text, fontsize=14)
                    y_offset += 25

                    img_width = pixmap.width
                    img_height = pixmap.height
                    max_img_width = new_page.rect.width - 2 * margin
                    max_img_height_on_current_page = new_page.rect.height - y_offset - margin
                    scale_x = max_img_width / img_width
                    scale_y = max_img_height_on_current_page / img_height
                    scale_factor = min(scale_x, scale_y)
                    scaled_img_width = img_width * scale_factor
                    scaled_img_height = img_height * scale_factor

                    if scaled_img_height > max_img_height_on_current_page:
                        new_page = output_doc.new_page(width=595, height=842)
                        y_offset = margin
                        new_page.insert_text((margin, y_offset), f"(Continued from Original Page {original_page_num})", fontsize=10)
                        y_offset += 20
                        max_img_height_on_current_page = new_page.rect.height - y_offset - margin
                        scale_y = max_img_height_on_current_page / img_height
                        scale_factor = min(scale_x, scale_y)
                        scaled_img_width = img_width * scale_factor
                        scaled_img_height = img_height * scale_factor

                    img_x = margin + (max_img_width - scaled_img_width) / 2
                    img_y = y_offset

                    target_rect = fitz.Rect(img_x, img_y, img_x + scaled_img_width, img_y + scaled_img_height)
                    new_page.insert_image(target_rect, pixmap=pixmap)
                    logging.info(f"Inserted full page screenshot for original page {original_page_num}")

            elif mode == 'cropped_highlight':
                y_offset = 50
                margin = 50
                max_page_width = 595
                max_page_height = 842
                current_page = output_doc.new_page(width=max_page_width, height=max_page_height)
                
                for highlight in highlights:
                    page_num = highlight['page']
                    rect_coords = highlight['rect']
                    rect = fitz.Rect(rect_coords)
                    
                    original_page = original_doc.load_page(page_num - 1)
                    matrix = fitz.Matrix(2, 2)
                    pixmap = original_page.get_pixmap(matrix=matrix, clip=rect)

                    # Calculate image dimensions and scaling
                    img_width = pixmap.width
                    img_height = pixmap.height
                    max_img_width = current_page.rect.width - 2 * margin
                    
                    scale_x = max_img_width / img_width
                    scale_factor = min(scale_x, 1.0) # Do not upscale
                    
                    scaled_img_width = img_width * scale_factor
                    scaled_img_height = img_height * scale_factor

                    # Check if the image will fit on the current page
                    if y_offset + scaled_img_height + margin > max_page_height:
                        current_page = output_doc.new_page(width=max_page_width, height=max_page_height)
                        y_offset = margin
                    
                    # Add title for the highlight
                    title_text = f"Page {page_num} Highlight:"
                    current_page.insert_text((margin, y_offset), title_text, fontsize=10)
                    y_offset += 15

                    # Insert the image
                    img_x = margin
                    img_y = y_offset
                    target_rect = fitz.Rect(img_x, img_y, img_x + scaled_img_width, img_y + scaled_img_height)
                    current_page.insert_image(target_rect, pixmap=pixmap)
                    
                    y_offset += scaled_img_height + 25 # Move down for the next highlight
                    logging.info(f"Inserted cropped highlight screenshot for page {page_num}")
            
            original_doc.close()

        except Exception as e:
            logging.exception(f"Error generating PDF from screenshots: {e}")
            output_doc = fitz.open()
            page = output_doc.new_page()
            page.insert_text((50, 50), "Error generating visual highlights. Please try again or check the original PDF.", fontsize=12)

    output_doc.save(output_pdf_path)
    output_doc.close()
    logging.info(f"Generated output PDF at: {output_pdf_path}")

def sanitize_text(text):
    """
    Removes invalid characters that can cause issues when writing to XML-based formats like DOCX.
    """
    # Remove control characters (except for common ones like tabs, newlines, etc.)
    # and any null bytes.
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

# --- NEW Function to detect complex text ---
def is_complex_text(text):
    """
    Checks if a string contains characters that are likely part of a formula or complex symbol.
    This is a simple heuristic and might need to be refined.
    """
    # Regex to detect common mathematical symbols, Greek letters, or non-ASCII characters that might
    # not render well as plain text in DOCX.
    complex_pattern = re.compile(r'[≤≥≠≈∞∫∑∏µΔΩ∂∇θλπρτφχψζηξςαβγδϵζηθικλμνξοπρσςτυφχψω]|[\u2200-\u22FF]|[\u0391-\u03C9]')
    return bool(complex_pattern.search(text))

# --- NEW Function to create an image from a highlight ---
def create_highlight_image(pdf_path, page_num, rect_coords):
    """
    Creates a high-resolution image of a specific highlighted area.
    Returns the path to the temporary image file.
    """
    temp_image_path = None
    try:
        doc = fitz.open(pdf_path)
        page = doc.load_page(page_num - 1)
        rect = fitz.Rect(rect_coords)
        
        # Increase the DPI for higher resolution
        matrix = fitz.Matrix(3, 3) 
        pixmap = page.get_pixmap(matrix=matrix, clip=rect)
        
        # Save the pixmap to a temporary file
        temp_image_path = tempfile.mktemp(suffix=".png", dir=app.config['UPLOAD_FOLDER'])
        pixmap.save(temp_image_path)
        
        doc.close()
        logging.info(f"Created temporary image for highlight: {temp_image_path}")
        return temp_image_path
    except Exception as e:
        logging.exception(f"Error creating image for highlight: {e}")
        return None

# --- Modified Function to Generate DOCX from Highlights ---
def generate_docx_from_highlights(original_pdf_path, highlights, output_docx_path):
    """
    Generates a new DOCX file containing all extracted highlights.
    If the text is complex, it inserts an image of the highlight.
    Otherwise, it inserts the text.
    """
    document = Document()
    document.add_heading('Extracted PDF Highlights', level=1)
    
    if not highlights:
        document.add_paragraph("No highlights found in the document.")
    else:
        for highlight in highlights:
            # Add a paragraph with the page number
            page_paragraph = document.add_paragraph()
            page_paragraph.add_run(f"Page {highlight['page']}:").bold = True
            
            # --- MODIFIED LOGIC ---
            text_paragraph = document.add_paragraph()
            sanitized_text = sanitize_text(highlight['text'])
            
            if is_complex_text(sanitized_text):
                logging.info(f"Complex text detected. Generating image for highlight on page {highlight['page']}.")
                image_path = create_highlight_image(original_pdf_path, highlight['page'], highlight['rect'])
                if image_path and os.path.exists(image_path):
                    try:
                        # Insert the image into the DOCX
                        document.add_picture(image_path, width=Inches(6)) # Adjust width as needed
                        logging.info("Successfully inserted image into DOCX.")
                    except Exception as e:
                        logging.error(f"Failed to insert image into DOCX: {e}")
                        text_paragraph.add_run(f"[[Image not available for complex text on Page {highlight['page']}]]")
                    finally:
                        # Clean up the temporary image file
                        os.remove(image_path)
                        logging.info(f"Cleaned up temporary image file: {image_path}")
                else:
                    logging.warning("Image creation failed, inserting sanitized text as fallback.")
                    text_paragraph.add_run(sanitized_text)
            else:
                text_paragraph.add_run(sanitized_text)
    
    document.save(output_docx_path)
    logging.info(f"Generated DOCX file at: {output_docx_path}")

# --- API Endpoint for PDF Upload and Processing ---
@app.route('/upload-pdf', methods=['POST'])
def upload_pdf():
    logging.info("Received request to /upload-pdf")

    if 'pdfFile' not in request.files:
        logging.warning("No 'pdfFile' part in request.")
        return jsonify({"error": "No file part in the request"}), 400

    file = request.files['pdfFile']
    extraction_mode = request.form.get('extractionMode', 'full_page')
    
    if file.filename == '':
        logging.warning("No selected file name provided.")
        return jsonify({"error": "No selected file"}), 400

    if not file.filename.lower().endswith('.pdf'):
        logging.warning(f"Invalid file type uploaded: {file.filename}")
        return jsonify({"error": "Invalid file type. Please upload a PDF document."}), 400

    if file:
        temp_input_pdf_path = None
        output_pdf_path = None
        output_docx_path = None
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf", dir=app.config['UPLOAD_FOLDER']) as temp_input_pdf:
                file.save(temp_input_pdf.name)
                temp_input_pdf_path = temp_input_pdf.name
            logging.info(f"Uploaded PDF saved temporarily to: {temp_input_pdf_path}")

            original_filename_base = os.path.splitext(file.filename)[0]
            
            # --- Process and generate PDF output ---
            output_pdf_id = uuid.uuid4().hex
            output_filename = f"highlights_{original_filename_base}_{output_pdf_id}.pdf"
            output_pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
            extracted_highlights = extract_highlights_from_pdf(temp_input_pdf_path)
            generate_pdf_from_highlight_screenshots(temp_input_pdf_path, extracted_highlights, output_pdf_path, extraction_mode)
            
            # --- Process and generate DOCX output ---
            output_docx_id = uuid.uuid4().hex
            output_docx_filename = f"highlights_text_{original_filename_base}_{output_docx_id}.docx"
            output_docx_path = os.path.join(app.config['UPLOAD_FOLDER'], output_docx_filename)
            generate_docx_from_highlights(temp_input_pdf_path, extracted_highlights, output_docx_path)

            pdf_download_url = f"/download-pdf/{output_filename}"
            docx_download_url = f"/download-docx/{output_docx_filename}"

            logging.info(f"Generated PDF available for download at: {pdf_download_url}")
            logging.info(f"Generated DOCX available for download at: {docx_download_url}")

            return jsonify({
                "highlights": extracted_highlights,
                "pdf_download_url": pdf_download_url,
                "docx_download_url": docx_download_url
            }), 200

        except Exception as e:
            logging.exception("An error occurred during PDF processing and generation.")
            return jsonify({"error": "Failed to process PDF. Please try again."}), 500
        finally:
            if temp_input_pdf_path and os.path.exists(temp_input_pdf_path):
                os.remove(temp_input_pdf_path)
                logging.info(f"Cleaned up input PDF: {temp_input_pdf_path}")
            pass

    logging.error("An unexpected flow occurred in /upload-pdf.")
    return jsonify({"error": "An unexpected error occurred"}), 500

# --- New Endpoint to Serve Generated PDF ---
@app.route('/download-pdf/<filename>', methods=['GET'])
def download_pdf(filename):
    logging.info(f"Received request to download PDF: {filename}")
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    if not os.path.exists(file_path):
        logging.warning(f"Requested file not found: {file_path}")
        return jsonify({"error": "File not found."}), 404

    try:
        response = send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True, mimetype='application/pdf')
        
        @response.call_on_close
        def cleanup():
            if os.path.exists(file_path):
                os.remove(file_path)
                logging.info(f"Cleaned up output PDF after download: {file_path}")
        
        return response
    except Exception as e:
        logging.exception(f"Error serving download file {filename}.")
        return jsonify({"error": "Could not serve the requested file."}), 500

# --- New Endpoint to Serve Generated DOCX ---
@app.route('/download-docx/<filename>', methods=['GET'])
def download_docx(filename):
    logging.info(f"Received request to download DOCX: {filename}")
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    
    if not os.path.exists(file_path):
        logging.warning(f"Requested file not found: {file_path}")
        return jsonify({"error": "File not found."}), 404
    
    try:
        response = send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        @response.call_on_close
        def cleanup():
            if os.path.exists(file_path):
                os.remove(file_path)
                logging.info(f"Cleaned up output DOCX after download: {file_path}")
        
        return response
    except Exception as e:
        logging.exception(f"Error serving download DOCX file {filename}.")
        return jsonify({"error": "Could not serve the requested file."}), 500

# --- Serve the frontend HTML directly from Flask for easier local testing ---
@app.route('/')
def serve_frontend():
    return send_from_directory(app.static_folder, 'index.html')

if __name__ == '__main__':
    static_dir = os.path.join(app.root_path, 'static')
    os.makedirs(static_dir, exist_ok=True)
    logging.info(f"Static directory for frontend: {static_dir}")
    app.run(host='0.0.0.0', port=8000, debug=False)